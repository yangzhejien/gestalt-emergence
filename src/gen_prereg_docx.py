# -*- coding: utf-8 -*-
"""Convert OSF_preregistration_2026-08-07.md -> .docx with clean typography.
Improvements over naive export:
  - LaTeX math ($...$ / $$...$$) converted to readable Unicode math (no raw '$').
  - Inline `code` spans preserved verbatim.
  - [Status] blockquote rendered as a shaded note box.
  - Justified body, styled headings, proper fonts (Times + SimSun for CJK),
    page margins, line spacing.
"""
import re, os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = r"D:\方程验证\OSF_preregistration_2026-08-07.md"
OUT = r"D:\方程验证\OSF_preregistration_2026-08-07_v2.docx"
REPORTS = r"D:\方程验证\reports"
LAT = "Consolas"
CJK = "SimSun"
HEAD_COLOR = RGBColor(0x1F, 0x38, 0x64)

# ---------------- inline code splitter ----------------
def split_code(text):
    """Return list of (segment, is_code)."""
    parts = text.split('`')
    return [(parts[i], i % 2 == 1) for i in range(len(parts))]

# ---------------- LaTeX -> Unicode math ----------------
SUB = {'0':'₀','1':'₁','2':'₂','3':'₃','4':'₄','5':'₅','6':'₆','7':'₇','8':'₈','9':'₉',
       'a':'ₐ','e':'ₑ','h':'ₕ','i':'ᵢ','j':'ⱼ','k':'ₖ','l':'ₗ','m':'ₘ','n':'ₙ','o':'ₒ','p':'ₚ','r':'ᵣ','s':'ₛ','t':'ₜ','u':'ᵤ','v':'ᵥ','x':'ₓ','+':'₊','-':'₋','=':'₌','(':'₍',')':'₎'}
SUP = {'0':'⁰','1':'¹','2':'²','3':'³','4':'⁴','5':'⁵','6':'⁶','7':'⁷','8':'⁸','9':'⁹',
       'a':'ᵃ','b':'ᵇ','c':'ᶜ','d':'ᵈ','e':'ᵉ','f':'ᶠ','g':'ᵍ','h':'ʰ','i':'ⁱ','j':'ʲ','k':'ᵏ','l':'ˡ','m':'ᵐ','n':'ⁿ','o':'ᵒ','p':'ᵖ','r':'ʳ','s':'ˢ','t':'ᵗ','u':'ᵘ','v':'ᵛ','x':'ˣ','+':'⁺','-':'⁻','=':'⁼','(':'⁽',')':'⁾'}

def to_sub(s):
    out = []; fb = False
    for ch in s:
        if ch in SUB: out.append(SUB[ch])
        else: out.append(ch); fb = True
    r = ''.join(out)
    return ('₍' + r + '₎') if (fb and len(s) > 1) else r

def to_sup(s):
    out = []; fb = False
    for ch in s:
        if ch in SUP: out.append(SUP[ch])
        else: out.append(ch); fb = True
    r = ''.join(out)
    return ('⁽' + r + '⁾') if (fb and len(s) > 1) else r

def _grab_brace(s, i):
    """s[i]=='{'; return (inner, index_after_close)."""
    depth = 0; j = i
    while j < len(s):
        if s[j] == '{': depth += 1
        elif s[j] == '}':
            depth -= 1
            if depth == 0:
                return s[i+1:j], j+1
        j += 1
    return s[i+1:], len(s)

def tex2uni(s):
    # spacing commands -> space
    s = re.sub(r'\\(quad|qquad|;|,|!| )', ' ', s)
    s = s.replace('~', ' ')
    # \text{...}
    def repl_text(m):
        inner, _ = _grab_brace(s, m.end()-1) if False else (None, 0)
        return ''
    # process \text{...}
    while True:
        m = re.search(r'\\text\s*\{', s)
        if not m: break
        inner, after = _grab_brace(s, m.end()-1)
        s = s[:m.start()] + inner + s[after:]
    # \frac{a}{b}
    while True:
        m = re.search(r'\\frac\s*\{', s)
        if not m: break
        a, after1 = _grab_brace(s, m.end()-1)
        m2 = re.search(r'\s*\{', s[after1:])
        b, after2 = _grab_brace(s[after1:], m2.start())
        s = s[:m.start()] + f'({a}/{b})' + s[after1+m2.start()+after2:]
    # \sqrt{...}
    while True:
        m = re.search(r'\\sqrt\s*\{', s)
        if not m: break
        inner, after = _grab_brace(s, m.end()-1)
        s = s[:m.start()] + '√(' + inner + ')' + s[after:]
    # \hat{...}
    while True:
        m = re.search(r'\\hat\s*\{', s)
        if not m: break
        inner, after = _grab_brace(s, m.end()-1)
        if inner and inner[0] in ('W',):
            hat = 'Ŵ' if inner[0] == 'W' else inner[0] + '̂'
        else:
            hat = (inner[0] + '̂' + inner[1:]) if inner else ''
        s = s[:m.start()] + hat + s[after:]
    # symbol commands
    sym = {
        r'\alpha':'α', r'\beta':'β', r'\lambda':'λ', r'\mu':'μ', r'\sigma':'σ',
        r'\sum':'Σ', r'\Sigma':'Σ', r'\times':'×', r'\cdot':'·', r'\leftrightarrow':'↔',
        r'\rightarrow':'→', r'\to':'→', r'\sim':'∼', r'\approx':'≈',
        r'\Rightarrow':'⇒', r'\dots':'…', r'\ell':'ℓ', r'\langle':'⟨', r'\rangle':'⟩',
        r'\nabla':'∇', r'\partial':'∂', r'\infty':'∞', r'\propto':'∝',
        r'\neq':'≠', r'\leq':'≤', r'\geq':'≥',
    }
    for k, v in sym.items():
        s = s.replace(k, v)
    # remove delimiter/style commands (keep following content)
    for cmd in [r'\\left', r'\\right', r'\\big', r'\\Big', r'\\bigg', r'\\Bigg',
                r'\\mathrm', r'\\mathit', r'\\mathbf', r'\\operatorname', r'\\displaystyle']:
        s = re.sub(cmd + r'\s*', '', s)
    # subscripts / superscripts (brace groups, then single char)
    while True:
        m = re.search(r'_\{', s)
        if not m: break
        inner, after = _grab_brace(s, m.end()-1)
        s = s[:m.start()] + to_sub(inner) + s[after:]
    while True:
        m = re.search(r'\^\{', s)
        if not m: break
        inner, after = _grab_brace(s, m.end()-1)
        s = s[:m.start()] + to_sup(inner) + s[after:]
    s = re.sub(r'_([A-Za-z0-9])', lambda m: to_sub(m.group(1)), s)
    s = re.sub(r'\^([A-Za-z0-9])', lambda m: to_sup(m.group(1)), s)
    # strip remaining braces and unknown backslash commands
    s = s.replace('{', '').replace('}', '')
    s = re.sub(r'\\[A-Za-z]+', '', s)
    s = s.replace('∥', '∥').strip()
    return s

# ---------------- run font helper ----------------
def set_run_font(run, eastasia=CJK, latin=LAT):
    run.font.name = latin
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), eastasia)
    rFonts.set(qn('w:ascii'), latin)
    rFonts.set(qn('w:hAnsi'), latin)

# ---------------- inline parsing (bold / italic / math / code) ----------------
INLINE = re.compile(r'(\*\*.+?\*\*|\*.+?\*|\$+[^\$]+\$+|\$[^\$]+\$)')

def add_inline(paragraph, text, math_italic=True):
    pos = 0
    for m in INLINE.finditer(text):
        if m.start() > pos:
            add_plain(paragraph, text[pos:m.start()])
        tok = m.group(0)
        if tok.startswith('**'):
            r = paragraph.add_run(tok[2:-2]); r.bold = True; set_run_font(r, latin=LAT, eastasia=CJK)
        elif tok.startswith('*'):
            r = paragraph.add_run(tok[1:-1]); r.italic = True; set_run_font(r, latin=LAT, eastasia=CJK)
        elif tok.startswith('$$'):
            r = paragraph.add_run(tex2uni(tok[2:-2])); r.italic = True; r.font.name = LAT; set_run_font(r, latin=LAT)
        else:  # $...$
            r = paragraph.add_run(tex2uni(tok[1:-1])); r.italic = True; r.font.name = LAT; set_run_font(r, latin=LAT)
        pos = m.end()
    if pos < len(text):
        add_plain(paragraph, text[pos:])

def add_plain(paragraph, text):
    for seg, is_code in split_code(text):
        if is_code:
            r = paragraph.add_run(seg); r.font.name = LAT; set_run_font(r, latin=LAT)
        else:
            add_inline_rich(paragraph, seg)

def add_inline_rich(paragraph, text):
    # handle code already split; this does bold/italic/math on non-code text
    pos = 0
    for m in INLINE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        tok = m.group(0)
        if tok.startswith('**'):
            r = paragraph.add_run(tok[2:-2]); r.bold = True; set_run_font(r, latin=LAT, eastasia=CJK)
        elif tok.startswith('*'):
            r = paragraph.add_run(tok[1:-1]); r.italic = True; set_run_font(r, latin=LAT, eastasia=CJK)
        elif tok.startswith('$$'):
            r = paragraph.add_run(tex2uni(tok[2:-2])); r.italic = True; r.font.name = LAT; set_run_font(r, latin=LAT)
        else:
            r = paragraph.add_run(tex2uni(tok[1:-1])); r.italic = True; r.font.name = LAT; set_run_font(r, latin=LAT)
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])

# ---------------- paragraph shading ----------------
def shade(paragraph, fill='F2F2F2'):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
    pPr.append(shd)

def list_level(line):
    m = re.match(r'^(\s*)', line)
    return len(m.group(1)) // 2

def main():
    lines = open(SRC, encoding='utf-8').read().split('\n')
    doc = Document()
    # base style
    normal = doc.styles['Normal']
    normal.font.name = 'Times New Roman'
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), CJK)
    pf = normal.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    # margins
    for s in doc.sections:
        s.left_margin = Inches(1.0); s.right_margin = Inches(1.0)
        s.top_margin = Inches(1.0); s.bottom_margin = Inches(1.0)
    # heading colors
    for hs in ['Heading 1', 'Heading 2', 'Heading 3', 'Title']:
        try:
            st = doc.styles[hs]
            st.font.color.rgb = HEAD_COLOR
            st.font.name = 'Times New Roman'
            st.element.rPr.rFonts.set(qn('w:eastAsia'), CJK)
        except KeyError:
            pass

    i = 0; n = len(lines)
    while i < n:
        line = lines[i]; stripped = line.strip()
        if stripped == '':
            i += 1; continue
        if stripped == '---':
            i += 1; continue
        mh = re.match(r'^(#{1,3})\s+(.*)$', stripped)
        if mh:
            lvl = len(mh.group(1)); txt = mh.group(2)
            if lvl == 1:
                p = doc.add_paragraph(txt, style='Title')
            else:
                p = doc.add_heading(level=lvl-1)
                add_inline(p, txt)
            i += 1; continue
        mi = re.match(r'^!\[(.*?)\]\((.*?)\)\s*$', stripped)
        if mi:
            alt, rel = mi.group(1), mi.group(2)
            base = os.path.splitext(os.path.basename(rel))[0]
            png = os.path.join(REPORTS, base + '.png')
            svg = os.path.join(REPORTS, base + '.svg')
            if os.path.exists(png):
                doc.add_picture(png, width=Inches(6.0))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = cap.add_run(alt); r.italic = True; r.font.size = Pt(9); set_run_font(r, latin=LAT)
            elif os.path.exists(svg):
                doc.add_paragraph(f"[图: {alt} — 见随附 {os.path.basename(rel)}]")
            else:
                doc.add_paragraph(f"[图: {alt}]")
            i += 1; continue
        if stripped.startswith('|') and i+1 < n and re.match(r'^\s*\|?[\s:\-|]+\|?\s*$', lines[i+1]):
            tbl_rows = []
            j = i
            while j < n and lines[j].strip().startswith('|'):
                tbl_rows.append([c.strip() for c in lines[j].strip().strip('|').split('|')])
                j += 1
            header = tbl_rows[0]; body = tbl_rows[2:]
            t = doc.add_table(rows=1, cols=len(header)); t.style = 'Light Grid Accent 1'
            for k, h in enumerate(header):
                add_inline(t.rows[0].cells[k].paragraphs[0], h)
            for row in body:
                cells = t.add_row().cells
                for k, c in enumerate(row):
                    add_inline(cells[k].paragraphs[0], c)
            i = j; continue
        if stripped.startswith('>'):
            content = stripped.lstrip('>').strip()
            if content.startswith('$$') and content.endswith('$$'):
                p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(tex2uni(content[2:-2])); r.italic = True; r.font.name = LAT
                r.font.size = Pt(12); set_run_font(r, latin=LAT)
                p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(8)
            elif content.startswith('[Status]'):
                p = doc.add_paragraph()
                rl = p.add_run('Status — '); rl.bold = True; rl.font.color.rgb = HEAD_COLOR; set_run_font(rl)
                add_plain(p, content[len('[Status]'):].strip())
                shade(p, 'EAF1FB')
                p.paragraph_format.left_indent = Inches(0.15); p.paragraph_format.right_indent = Inches(0.15)
            else:
                p = doc.add_paragraph(); add_inline(p, content); p.italic = True
            i += 1; continue
        if re.match(r'^\s*([-\*]|\d+\.)\s+', stripped):
            while i < n and re.match(r'^\s*([-\*]|\d+\.)\s+', lines[i].strip()):
                ln = lines[i].strip(); lvl = list_level(lines[i])
                mm = re.match(r'^(\s*)([-\*]|\d+\.)\s+(.*)$', ln)
                lead, marker, txt = mm.group(1), mm.group(2), mm.group(3)
                if marker in ('-', '*'):
                    style = 'List Bullet' if lvl == 0 else f'List Bullet {lvl+1}'
                else:
                    style = 'List Number' if lvl == 0 else f'List Number {lvl+1}'
                p = doc.add_paragraph(style=style)
                add_inline(p, txt)
                i += 1
            continue
        p = doc.add_paragraph(); add_inline(p, stripped)
        i += 1

    doc.save(OUT)
    print("WROTE", OUT, os.path.getsize(OUT), "bytes")

if __name__ == '__main__':
    main()
