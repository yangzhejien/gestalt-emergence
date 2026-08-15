# -*- coding: utf-8 -*-
"""生成「格式塔方程验证 — 可复现性环境与复现指南」正式 Word 文档。"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

OUT = r"D:\方程验证\可复现包_2026-08-11\环境说明与复现指南.docx"

doc = Document()

# 默认字体（含中文）
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")

def set_cjk(run, font="SimSun"):
    run.font.name = "SimSun"
    r = run._element
    r.rPr.rFonts.set(qn("w:eastAsia"), font)

def heading(text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_cjk(run)
    return h

def para(text, bold=False, size=11, align=None, color=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    set_cjk(r)
    return p

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    r = p.add_run(text)
    r.font.name = "Courier New"
    r.font.size = Pt(9)
    # 等宽字体中英都设
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
    return p

def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    set_cjk(r)
    return p

# ---------------- 标题 ----------------
title = doc.add_heading("格式塔方程验证：可复现性环境与复现指南", level=0)
for run in title.runs:
    set_cjk(run)
para("技术文档 · Reproducibility & Environment Specification", size=10,
     align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x66, 0x66, 0x66))

# ---------------- 摘要 ----------------
heading("摘要", 1)
para("本文档规定「格式塔方程」主基准涌现实验的软件环境与复现规程，用于在独立计算"
     "平台上验证：分层多模型协作架构（L3 专家集群 → L2 副脑 → L1 主脑）的集体准确率"
     "是否稳定超越其最强单体组件（7B 模型）。文档涵盖操作系统与运行时版本、软件依赖、"
     "模型权重获取、数据集结构、复现命令、预期结果对照、已知环境约束，以及复现贡献记录规范。"
     "本包不随附模型权重，复现者须通过 Ollama 自行拉取指定版本的 Qwen2.5 系列模型。")

# ---------------- 1 概述 ----------------
heading("1. 概述与目的", 1)
para("本项目提出一种带符号的涌现度量方程，并主张：在适当的跨层连接强度（Ẇ）下，"
     "分层协作的集体性能可超线性地超过其最强单体。为验证该主张，主基准采用程序化生成的"
     "多步推理 MCQ（模板出题、答案由代码计算、随机种子固定），在 temp=0 下由本地模型推理，"
     "从而保证结果确定性、可复现。本包的目标，是让第三方研究者在自有机器上独立重跑关键实验，"
     "确认涌现现象的稳健性（环境/可移植性维度）。")

# ---------------- 2 系统环境 ----------------
heading("2. 系统环境规范", 1)
heading("2.1 操作系统与运行时", 2)
bullet("操作系统：Windows 10/11、Linux、macOS 均可；原实验在 Windows 11 完成。")
bullet("Python：3.13.12（3.10 及以上版本应兼容）。")
bullet("Ollama：0.32.8（或其他近期稳定版本；模型推理后端）。")
bullet("本机访问 Ollama 须使用 127.0.0.1，而非 localhost（见 §7）。")

heading("2.2 硬件建议", 2)
bullet("内存：≥ 8 GB 可运行 7B CPU 推理；推荐 16 GB 及以上以避免交换。")
bullet("存储：模型权重约 5.5 GB（1.5B+3B+7B），加数据集与结果，预留 ≥ 10 GB。")
bullet("GPU：可选。GPU 仅加速推理，不改变最终准确率；本包默认按 CPU 串行设计。")

# ---------------- 3 依赖 ----------------
heading("3. 软件依赖与安装", 1)
heading("3.1 Python 依赖", 2)
para("运行复现仅需 numpy；其余依赖均为 Python 标准库（urllib、json、re、argparse 等），"
     "无需额外安装。文档/报告生成所用的 python-docx 不属于复现运行依赖。")
code_block("pip install -r requirements.txt")
para("requirements.txt 内容：numpy>=1.24。")

heading("3.2 Ollama 与模型权重获取", 2)
para("模型权重由 Ollama 官方仓库提供，出于许可与体积考量不随包分发，每台机器须自行拉取"
     "以下三个标签（架构在 verify_stage2.py 中固定）：")
bullet("l3 专家集群：qwen2.5:1.5b（按集群大小 k 重复调用）")
bullet("聚合层 / L2 副脑：qwen2.5:3b")
bullet("L1 主脑：qwen2.5:7b")
bullet("验证层：qwen2.5:1.5b")
code_block("ollama pull qwen2.5:1.5b\nollama pull qwen2.5:3b\nollama pull qwen2.5:7b")

# ---------------- 4 数据集 ----------------
heading("4. 数据集说明", 1)
heading("4.1 主基准 mcq_medium_clean.jsonl", 2)
para("500 道程序化生成的多步算术/逻辑 MCQ，每行格式为 "
     "{question, A, B, C, D, answer}。答案由出题代码直接计算，"
     "因此评测不依赖任何外部标注，且 temp=0 下模型输出确定性。这是 E3（k=5）与"
     "条件C（k=3）复现的主题库。")
heading("4.2 子集与对照集 / 生成器", 2)
bullet("mcq_medium.jsonl：主基准母集 500 题，供编排器抽取子集。")
bullet("mcq_medium_sub200_s20260811.jsonl：固定随机种子（20260811）抽取的 200 题子集，用于验证涌现非特定题集偶然。")
bullet("ceval_bandok_clean.jsonl：C-Eval 中文硬知识基准（外部交叉验证，反例）。")
bullet("ceval_en.jsonl：C-Eval 同题英文翻译（语言对照）。")
bullet("mcq_medium_orig100.jsonl：主基准最初的 100 道种子题（expand_to_500.py 的输入）。")
bullet("mcq_small.jsonl：早期开发用 100 题小基准（DEFAULT_CFG 的历史默认值，现已改为主基准）。")
para("基准再生（审计/透明用，复现者直接用上面的 .jsonl 即可，无需重跑）：", size=10)
bullet("code/expand_to_500.py（SEED=20260805，确定性）：在 orig100 基础上程序化生成 400 道同分布题 → 写回 mcq_medium.jsonl。")
bullet("code/clean_benchmark.py（SEED=20260804，确定性）：清洗 mcq_medium.jsonl（修占位符、选项重排、整体乱序）→ mcq_medium_clean.jsonl。")
para("注意：code/gen_bench.py 生成的是独立的 mcq_hard100（高难题库），并非主基准，属辅助脚本。", size=10)

# ---------------- 5 复现规程 ----------------
heading("5. 复现规程", 1)
heading("5.1 一键编排模式（推荐）", 2)
para("编辑 code/run_light_repro.py 顶部三个变量为本机绝对路径：")
code_block("OUT  = Path(r'你的结果输出目录')\nROOT = Path(r'本包根目录')\nPY   = r'你的 python 解释器路径'")
para("随后运行，将串行执行两项复现（自带 Ollama 探活与题级断点续跑）：")
code_block("python code/run_light_repro.py")
bullet("① E3 k=5 同配置重跑（N=500, temp=0）→ 验证原值 0.942 重现")
bullet("② 换样本子集条件C（k=3, N=200, temp=0）→ 验证涌现非题集偶然")

heading("5.2 手动单跑模式（更可控）", 2)
para("直接调用主运行器，显式指定参数，无需修改脚本（务必用 --out 指定你自己的结果目录）：")
code_block("python code/verify_stage2.py --k 5 --n 500 --conn-w 1.0 \\\n"
           "  --benchmark benchmark/mcq_medium_clean.jsonl --temperature 0.0 \\\n"
           "  --out results --live results/my_E3_k5.json")
para("中断后可重启同命令续跑：进度保存在 <--out 目录>/tiers/<live_stem>_cw*.jsonl。")

heading("5.3 关键参数语义", 2)
bullet("--k：L3 专家集群大小（主基准取 5）。")
bullet("--n：题目数量。")
bullet("--conn-w（Ẇ）：跨层连接强度，默认 1.0（本包定点复现）。")
bullet("--temperature：采样温度；主基准为 0.0（确定性）。")
bullet("--benchmark：题库路径；缺省时使用脚本内默认配置（现为主基准 mcq_medium_clean.jsonl）。")
bullet("--out：结果输出目录（live json / tiers / 报告）；缺省为本机 gestalt_live，**复现者务必用 --out 指定自己的目录**。")
bullet("--live：结果输出 json 文件名（位于 --out 目录下）。")

# ---------------- 6 预期结果 ----------------
heading("6. 预期结果与原值对照", 1)
para("下表为原实验已验证结果，复现者应得到一致数值（主基准确定性，见 §7）。")

tbl = doc.add_table(rows=1, cols=4)
tbl.style = "Table Grid"
hdr = tbl.rows[0].cells
for i, t in enumerate(["实验", "集体准确率", "最强单体(7B)", "状态"]):
    hdr[i].text = t
    for p in hdr[i].paragraphs:
        for r in p.runs:
            r.bold = True
            set_cjk(r)
rows = [
    ("E3 k=5（N=500）", "≈ 0.942", "0.836", "已验证"),
    ("条件C 复现（N=500）", "0.924", "0.838", "已验证"),
    ("换子集条件C（N=200）", "待复现", "—", "验证非题集偶然"),
    ("C-Eval k=1（中文）", "0.66", "0.785", "反例·无涌现"),
    ("容量对照 L3=3B", "0.72", "0.785", "反例·无涌现"),
    ("语言对照（英文）", "0.585", "0.675", "反例·无涌现"),
]
for exp, coll, single, st in rows:
    c = tbl.add_row().cells
    c[0].text, c[1].text, c[2].text, c[3].text = exp, coll, single, st
    for cell in c:
        for p in cell.paragraphs:
            for r in p.runs:
                set_cjk(r)

para("说明：主基准（前三项）呈现协作集体显著超越最强单体（约 +10 个百分点），构成核心"
     "涌现主张；C-Eval 三项为外部反例，用于论文中正面解释「互补结构」机制（硬知识题上"
     "7B 占优、小模型仅灌噪声 → 聚合回归弱层均值），属既定分析，复现者了解即可。", size=10)

# ---------------- 7 约束 ----------------
heading("7. 已知约束与环境稳健性说明", 1)
bullet("地址解析：必须使用 127.0.0.1，localhost 在部分环境被解析为 IPv6 ::1 会导致 Ollama 请求卡死。")
bullet("并发限制：CPU 推理并发 ≤ 2 客户端，多开会互相饿死，必须串行执行。")
bullet("确定性：主基准 temp=0 + 答案代码算 → 模型输出确定性，同配置多机结果应接近逐字节相同。")
bullet("多机复现的科学含义：验证的是「环境/可移植性稳健性」，而非统计显著性（后者靠换子集/换模型家族实现）。")
bullet("后台任务：长时运行建议以后台进程启动并监控日志；脚本支持自动重启续跑。")
bullet("solo 基线缓存陷阱：单模型基线（solo_checkpoint.json，位于 --out 目录）仅以 (题库路径, k, n) 为键缓存，不含模型/种子。换模型后若复用同一 --out 目录且同 k/n，会读到旧基线 → 结果错误。换模型或换配置前，删除 <--out>/solo_checkpoint.json。")

# ---------------- 8 文件清单 ----------------
heading("8. 文件清单", 1)
bullet("code/：verify_stage2.py（主运行器）、run_experiments.sh（★全套启动命令总表）、run_light_repro.py（作者一键编排）、run_controls.py、run_n500_clean.sh、run_n500_demote.sh、gen_bench.py、gen_ceval_en.py、expand_to_500.py、clean_benchmark.py、view_results.py、cfg_ceval_ctrl_cap.json")
bullet("benchmark/：mcq_medium_clean.jsonl（主基准 500 题）、mcq_medium.jsonl、mcq_medium_sub200_*.jsonl、mcq_medium_orig100.jsonl、mcq_small.jsonl、ceval_*.jsonl")
bullet("results_reference/：原实验已验证结果快照（stage2_E3_k5.json 等 7 个）")
bullet("requirements.txt、README.md、本 Word 文档")

# ---------------- 9 贡献记录 ----------------
heading("9. 复现贡献记录规范", 1)
para("本实验采用多人独立复现 + 共同修正的工作模式。每位复现者须如实记录其实际贡献，"
     "作为论文作者署名（author contributions，CRediT 分类）的依据。**仅记录实际完成"
     "复现或修正代码/分析的贡献者；无实质贡献者不挂名**，以杜绝 gift authorship。")
ctbl = doc.add_table(rows=1, cols=5)
ctbl.style = "Table Grid"
ch = ctbl.rows[0].cells
for i, t in enumerate(["姓名", "机器规格", "复现了哪项", "提出/修正了什么", "日期"]):
    ch[i].text = t
    for p in ch[i].paragraphs:
        for r in p.runs:
            r.bold = True
            set_cjk(r)
for _ in range(6):
    c = ctbl.add_row().cells
    for cell in c:
        for p in cell.paragraphs:
            for r in p.runs:
                set_cjk(r)

# ---------------- 10 模型连接 ----------------
heading("10. 模型连接机制与扩展", 1)
para("复现脚本通过本地 Ollama 的 HTTP 接口调用模型，不依赖任何第三方 SDK。核心机制位于 "
     "verify_stage2.py 的 generate() 函数：向 http://127.0.0.1:11434/api/generate 发送 "
     "POST 请求（Python 标准库 urllib.request），返回的响应文本即模型回答。Ollama 服务须"
     "在本机常驻（ollama serve 或桌面客户端）。")
bullet("连接地址：cfg['ollama_url']，默认 http://127.0.0.1:11434/api/generate（必须用 127.0.0.1，见 §7）。")
bullet("分层架构固定写在 cfg 的 models 字典：l3 专家集群=qwen2.5:1.5b（按集群大小 k 重复调用）、agg/L2 副脑=qwen2.5:3b、L1 主脑/编排=qwen2.5:7b、验证层=qwen2.5:1.5b。")
para("如何添加或更换模型：")
bullet("① 拉取新标签：ollama pull <新模型标签>，例如 ollama pull qwen2.5:14b。")
bullet("② 改架构：编辑 verify_stage2.py 顶部 cfg 的 models 字典（或 run_controls.py 所用的 cfg_ceval_ctrl_cap.json），把目标层指向新标签。")
bullet("③ 改连接主机/端口：编辑 cfg['ollama_url']；指向远端 Ollama 时须保证该端口可达，并已配置 CORS/代理。")
para("自检 Ollama 是否就绪：脚本内置探活（轮询 /api/tags）；也可手动执行 "
     "curl http://127.0.0.1:11434/api/tags 验证返回 200。", size=10)

# ---------------- 11 数据查看 ----------------
heading("11. 结果数据查看与解读", 1)
para("复现的主交付物是单个 JSON 文件（--live 指定，例如 results/my_E3_k5.json），"
     "除脚本自带断点续跑的 tiers/<stem>_cw*.jsonl 外，该 JSON 即核心结果。包内已提供解析脚本 "
     "code/view_results.py，直接打印关键指标，无需手动翻阅字段：")
code_block("python code/view_results.py results/my_E3_k5.json")
para("JSON 顶层关键字段含义：")
bullet("status / phase：运行态（running / done）与当前阶段（solo / pipeline / fit）。")
bullet("models：各层所用模型标签。")
bullet("n_questions：题目数。")
bullet("node_acc：每个节点的单独准确率列表（顺序为 k 个 L3 → L2 → L1）。")
bullet("best_single：最强单体准确率（主基准即 7B 节点值）。")
bullet("committee0：纯多数投票基线（构造对照，用于反制“只是平均模型/退化投票”指控）。")
bullet("collective_acc：协作集体准确率，键为 Ẇ 档（如 cw1.00），值为集体 acc。")
bullet("points：每个 Ẇ 档一行，含 G = collective − best_single（经验涌现锚）；superlinear 是否为真、judgement 文本结论。")
para("复现者应核对：collective_acc[cw1.00] 是否 ≈ 0.942（E3 k=5）、G > 0（集体超单体），"
     "并与 results_reference/ 原值快照及本文 §6 对照表比对一致。", size=10)

# ---------------- 12 参数总表 ----------------
heading("12. 实验参数总表（架构与全部参数）", 1)
para("下表汇总全部实验的架构与参数，确保复现者无需阅读源码即可完整配置，杜绝“靠猜”。")
ptbl = doc.add_table(rows=1, cols=3)
ptbl.style = "Table Grid"
ph = ptbl.rows[0].cells
for i, t in enumerate(["参数 / 项", "取值", "说明"]):
    ph[i].text = t
    for p in ph[i].paragraphs:
        for r in p.runs:
            r.bold = True; set_cjk(r)
prows = [
    ("分层架构", "L3 集群(×k, qwen2.5:1.5b) → 聚合层(3b) → L2 副脑(3b) → L1 主脑(7b)，每层旁挂验证层(1.5b)", "非投票的协作合成；聚合层输出「简报 + 残留」双通道"),
    ("L3 专家数 k", "主基准: 1(E1) / 3(条件C) / 5(E3) / 7(E3)；对照: 1", "取 cfg 中前 k 个 persona（共 7 个定义）"),
    ("连接强度 Ẇ (--conn-w)", "1.0（定点复现）", "0=只信聚合基；1=充分剥削跨层残留/分歧"),
    ("采样温度", "0.0（确定性）", "temp=0 下模型输出确定；--seed 仅在 temp>0 时才引入可复现随机性"),
    ("题库 (--benchmark)", "mcq_medium_clean.jsonl(主) / ceval_bandok_clean.jsonl / ceval_en.jsonl", "缺省=主基准；路径相对包根（benchmark/）"),
    ("题数 (--n)", "500（主基准）/ 200（C-Eval 系列）", "缺省=使用题库全部"),
    ("评测方式", "提取末选项字母，与 q['answer'] 精确比对", "答案由出题代码计算，不依赖外部标注 → temp=0 确定性"),
    ("涌现判据", "beats_best(集体>最强单体) 且 峰值落在 Ẇ>0", "若峰值在 Ẇ=0（加性平台）即便数值超单体也不判涌现"),
    ("消融模式 (--ablation)", "none（默认）/ demote（7B 降级进 L3 集群）/ remove（移除 7B）", "仅用于附验 D，主基准用 none"),
    ("solo 基线缓存", "solo_checkpoint.json（键=(题库, k, n)）", "换模型/换配置前须删除，否则读到旧基线（见 §7）"),
]
for a, b, c in prows:
    cells = ptbl.add_row().cells
    cells[0].text, cells[1].text, cells[2].text = a, b, c
    for cell in cells:
        for p in cell.paragraphs:
            for r in p.runs:
                set_cjk(r)

# ---------------- 13 命令总表 ----------------
heading("13. 启动命令总表", 1)
para("全部实验的精确启动命令见 code/run_experiments.sh（在本包根目录执行 "
     "bash code/run_experiments.sh 即可串行跑完全部）。每条命令均显式传 --out 指向你自己的"
     "结果目录。下表给出各实验的关键参数与对应的原值快照：")
ct2 = doc.add_table(rows=1, cols=4)
ct2.style = "Table Grid"
ch2 = ct2.rows[0].cells
for i, t in enumerate(["实验", "关键参数 (--live / --k / --n / --benchmark)", "原值快照", "状态"]):
    ch2[i].text = t
    for p in ch2[i].paragraphs:
        for r in p.runs:
            r.bold = True; set_cjk(r)
cmdrows = [
    ("E1 k=1", "stage2_E1_k1.json / k1 / n500 / 主基准", "stage2_E1_k1.json", "已验证"),
    ("条件C k=3", "stage2_condC_k3.json / k3 / n500 / 主基准", "condC_clean_relive.json(0.924)", "已验证"),
    ("E3 k=5", "stage2_E3_k5.json / k5 / n500 / 主基准", "stage2_E3_k5.json(0.942)", "已验证"),
    ("E3 k=7", "stage2_E3_k7.json / k7 / n500 / 主基准", "(本包未含, 见论文)", "已验证"),
    ("C-Eval k=1", "stage2_ceval_k1.json / k1 / n200 / ceval_bandok_clean", "stage2_ceval_k1.json", "反例"),
    ("C-Eval k=3", "stage2_ceval_k3.json / k3 / n200 / ceval_bandok_clean", "stage2_ceval_k3.json", "反例"),
    ("容量对照", "stage2_ceval_ctrl_cap_k1.json / k1 / n200 / 经 --config(3B集群)", "stage2_ceval_ctrl_cap_k1.json", "反例"),
    ("语言对照", "stage2_ceval_ctrl_lang_k1.json / k1 / n200 / ceval_en", "stage2_ceval_ctrl_lang_k1.json", "反例"),
]
for exp, params, snap, st in cmdrows:
    cells = ct2.add_row().cells
    cells[0].text, cells[1].text, cells[2].text, cells[3].text = exp, params, snap, st
    for cell in cells:
        for p in cell.paragraphs:
            for r in p.runs:
                set_cjk(r)
para("注：E3 k=7 原值快照未纳入本包（归档体积考虑），其数值与结论以论文正文为准；"
     "其余 7 个快照均在 results_reference/ 中，复现后逐条用 code/view_results.py 核对。", size=10)

# ---------------- 附录 ----------------
heading("附录：命令速查", 1)
code_block("# 安装依赖\npip install -r requirements.txt\n\n"
           "# 拉取模型\nollama pull qwen2.5:1.5b qwen2.5:3b qwen2.5:7b\n\n"
           "# 一键复现（先改 run_light_repro.py 三变量）\npython code/run_light_repro.py\n\n"
           "# 手动单跑 E3 k=5\npython code/verify_stage2.py --k 5 --n 500 --conn-w 1.0 "
           "--benchmark benchmark/mcq_medium_clean.jsonl --temperature 0.0 --out results --live results/my_E3_k5.json\n\n"
           "# 全套串行复现(推荐)：\nbash code/run_experiments.sh")

doc.save(OUT)
print("saved:", OUT)
